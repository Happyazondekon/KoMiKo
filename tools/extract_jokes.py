import docx
import json
import os

def extract():
    docx_path = r'C:/Users/ELITEBOOK/StudioProjects/Komiko/assets/Compilations de blagues_Projet Komiko.docx'
    json_path = r'C:/Users/ELITEBOOK/StudioProjects/Komiko/assets/jokes_import.json'

    if not os.path.exists(docx_path):
        print(f"Error: {docx_path} not found.")
        return

    doc = docx.Document(docx_path)
    jokes = []
    current_cat = 'Général'

    target_cats = [
        'Animaux', 'Belges', 'Blondes', 'Toto', 'Informatique',
        'Management', 'Médecine', 'Sport', 'Histoires US', 'Proverbes'
    ]

    print(f"Total paragraphs: {len(doc.paragraphs)}")

    current_joke_lines = []

    # Start after the table of contents (approx paragraph 730)
    for p in doc.paragraphs[730:]:
        text = p.text.strip()

        # Detect category change (usually short bold/header text)
        is_cat_header = False
        for c in target_cats:
            if c.lower() == text.lower() or (c.lower() in text.lower() and len(text) < 25):
                # Before changing category, save current joke
                if current_joke_lines:
                    jokes.append({
                        "category": current_cat,
                        "contentFr": "\n".join(current_joke_lines),
                        "punchlineFr": None,
                        "contentEn": None,
                        "punchlineEn": None
                    })
                    current_joke_lines = []

                current_cat = c
                is_cat_header = True
                break

        if is_cat_header:
            continue

        if not text:
            # Empty line usually separates jokes
            if current_joke_lines:
                # Merge lines and save
                content = "\n".join(current_joke_lines)
                if len(content) > 30: # Ignore very short artifacts
                    jokes.append({
                        "category": current_cat,
                        "contentFr": content,
                        "punchlineFr": None,
                        "contentEn": None,
                        "punchlineEn": None
                    })
                current_joke_lines = []
            continue

        # Accumulate lines for a joke
        current_joke_lines.append(text)

    # Final joke
    if current_joke_lines:
        jokes.append({
            "category": current_cat,
            "contentFr": "\n".join(current_joke_lines),
            "punchlineFr": None,
            "contentEn": None,
            "punchlineEn": None
        })

    print(f"Successfully extracted {len(jokes)} structured jokes.")

    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(jokes, f, indent=2, ensure_ascii=False)

    print(f"Saved to {json_path}")

if __name__ == "__main__":
    extract()
